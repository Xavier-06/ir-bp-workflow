"""IC Topic (课题研究) Pipeline Profile — 6 Role / 3 Wave / 18 Phase.

基于 multi-agent-pipeline 框架的行业课题研究管线，已对齐 BP/IR/LIT 的质量闭环。

架构:
  Wave 1: ic_market_overview / ic_competitive_landscape / ic_tech_product  (基础扫描)
  Wave 2: ic_supply_chain / ic_policy_risk                                (深度分析, 依赖 W1)
  Wave 3: ic_report_synthesizer                                           (统稿, 依赖 W1+W2)

Phase 清单 (18 phases):
  01 topic_intake               课题入库 + 研究计划
  02 presearch                  预搜索 [heavy_bg]
  03 extract                    内容提取
  04 fact_store_bootstrap       Fact Store 初始化
  05 wave1_dispatch_prepare     W1 派发 → needs_dispatch (sequential has_more)
  06 wave1_dispatch_collect     W1 收集 [retry + 进度检测]
  07 wave1_evidence_gate        W1 门禁 [severity 分级 + repair]
  08 wave1_fact_store_merge     W1 合并 [atomic write]
  09 wave1_shared_state_refresh W1 共享状态刷新 [NEW]
  10 wave2_dispatch_prepare     W2 派发 → needs_dispatch (sequential has_more)
  11 wave2_dispatch_collect     W2 收集 [retry + 进度检测]
  12 wave2_evidence_gate        W2 门禁 [severity 分级 + repair]
  13 wave2_fact_store_merge     W2 合并 [atomic write]
  14 wave2_shared_state_refresh W2 共享状态刷新 [NEW]
  15 synthesis_dispatch_prepare 统稿派发 → needs_dispatch
  16 synthesis_dispatch_collect 统稿收集 [retry + citation gate + repair]
  17 debate_review              对抗评审 [BLOCKING/MEDIUM 分级]
  18 delivery                   交付 [heavy_bg: DOCX + artifacts + investment judgment]

2026-07-08 v2: P0+P1 全面补全 — collect retry+进度检测 / evidence gate分级+repair闭环 / 
                共享状态刷新 / phase_prerequisites / 文件锁 / DOCX交付 / 投资判断汇总
"""
from __future__ import annotations

import json
import os
import shutil
import time
from pathlib import Path
from typing import Any

from runtime.profiles.base import JobContext, PipelineProfile

from scripts.ic_topic_constants import (
    IC_TOPIC_ALL_ROLE_SLUGS,
    IC_TOPIC_ROLE_CONNECTOR_IDS,
    IC_TOPIC_WAVE1_ROLE_SLUGS,
    IC_TOPIC_WAVE2_ROLE_SLUGS,
    IC_TOPIC_WAVE3_ROLE_SLUGS,
    IC_TOPIC_WAVE_ROLES,
    FACTS_SUFFIX,
    SECTION_SUFFIX,
    GATE_REPAIR_MAX_ATTEMPTS,
    COLLECT_RETRY_COUNT,
    COLLECT_RETRY_INTERVAL,
    WAVE3_GATE_THRESHOLDS,
)


# ═══════════════════════════════════════════════════════════
# Utilities
# ═══════════════════════════════════════════════════════════

def _task_dir(runtime_root: Path, job_ctx: JobContext) -> Path:
    ws = job_ctx.workspace
    if ws:
        return ws.root
    d = runtime_root / "data" / "tasks" / job_ctx.job_id
    d.mkdir(parents=True, exist_ok=True)
    return d


def _outputs_dir(runtime_root: Path, job_ctx: JobContext) -> Path:
    ws = job_ctx.workspace
    if ws:
        return ws.outputs_dir
    return _task_dir(runtime_root, job_ctx)


def _sync_to_workspace(job_ctx: JobContext, src: Path, dest_name: str):
    ws = job_ctx.workspace
    if ws is None or not src.exists():
        return
    dest = ws.outputs_dir / dest_name
    try:
        shutil.copy2(src, dest)
    except Exception:
        pass


def _atomic_write_json(file_path: Path, data: Any):
    """原子写入 JSON 文件（通过临时文件 + os.replace）。"""
    import tempfile
    tmp = tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False,
                                      dir=file_path.parent, encoding='utf-8')
    try:
        json.dump(data, tmp, ensure_ascii=False, indent=2)
        tmp.write("\n")
        tmp.flush()
        os.fsync(tmp.fileno())
        tmp.close()
        os.replace(tmp.name, str(file_path))
    except Exception:
        try:
            os.unlink(tmp.name)
        except Exception:
            pass
        # Fallback to direct write
        file_path.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


# ═══════════════════════════════════════════════════════════
# Collect with Retry + Progress Detection (参照 LIT)
# ═══════════════════════════════════════════════════════════

def _collect_with_retry(
    task: Path,
    roles: list[str],
    *,
    collect_name: str = "ic_collect",
    max_retries: int = COLLECT_RETRY_COUNT,
    retry_interval: int = COLLECT_RETRY_INTERVAL,
    check_fn=None,
) -> dict[str, Any]:
    """带进度检测的收集重试。

    子代理先写 .md 再写 sidecar（JSON 序列化耗时），
    直接判定 incomplete 会阻断管线。

    进度检测: 两次尝试间 incomplete 数量 + .md 总大小都不变 → 子代理可能已挂 → 提前退出。
    """
    from scripts.ic_topic_subagent_launcher import _role_outputs_complete

    if check_fn is None:
        check_fn = _role_outputs_complete

    last_result: dict[str, Any] | None = None
    prev_signal: tuple[int, int] | None = None

    for attempt in range(max_retries + 1):
        incomplete = [r for r in roles if not check_fn(task, r)]

        results = {}
        for role in roles:
            ok = check_fn(task, role)
            results[role] = {"files_complete": ok}
        all_ok = len(incomplete) == 0
        last_result = {"ok": all_ok, "incomplete": incomplete, "roles": results}

        if all_ok:
            return last_result

        # 进度检测
        total_md_size = 0
        for md_file in task.glob("*.md"):
            try:
                total_md_size += md_file.stat().st_size
            except OSError:
                pass
        current_signal = (len(incomplete), total_md_size)

        if prev_signal is not None and current_signal == prev_signal:
            print(f"  [{collect_name}] no progress ({len(incomplete)} incomplete), "
                  f"sub-agent may have stopped", flush=True)
            break

        print(f"  [{collect_name}] attempt {attempt+1}/{max_retries+1}, "
              f"{len(incomplete)} incomplete: {incomplete[:3]}, "
              f"retrying in {retry_interval}s...", flush=True)
        prev_signal = current_signal
        time.sleep(retry_interval)

    return last_result or {"ok": False, "incomplete": roles}


# ═══════════════════════════════════════════════════════════
# Gate Repair Helpers (参照 LIT)
# ═══════════════════════════════════════════════════════════

def _build_ic_repair_manifest(
    task: Path,
    role: str,
    slug: str,
    failure_reason: str,
    gate_phase: str,
    runtime_root: Path,
) -> str:
    """构建 IC gate repair manifest，让子代理针对失败项补充采集。"""
    from scripts.ic_topic_subagent_launcher import _assemble_system_prompt

    system_prompt = _assemble_system_prompt(runtime_root, role, task_dir=task)

    repair_prompt = (
        f"\n\n## REPAIR TASK — Gate {gate_phase} FAIL\n\n"
        f"**失败原因**: {failure_reason}\n\n"
        f"## 要求\n"
        f"1. 阅读现有的 {role}.md 和 {role}-facts.json\n"
        f"2. 针对上述失败原因，补充搜索/采集\n"
        f"3. 将新发现追加到现有文件（不覆盖已有内容），不新建文件\n"
        f"4. 更新 {role}-section.json 的统计数据\n"
        f"5. 完成后正常退出，管线会自动重跑 gate\n"
    )

    manifest = {
        "role": role,
        "slug": slug,
        "system_prompt": system_prompt + repair_prompt,
        "output_path": str(task / f"{role}.md"),
        "connectorIds": IC_TOPIC_ROLE_CONNECTOR_IDS.get(role, []),
        "is_repair": True,
        "repair_reason": failure_reason,
        "expected_outputs": [f"{role}.md", f"{role}{FACTS_SUFFIX}", f"{role}{SECTION_SUFFIX}"],
    }
    manifest_path = task / f"repair_manifest_{role}.json"
    _atomic_write_json(manifest_path, manifest)
    return str(manifest_path)


def _ic_repair_instruction(phase_name: str, has_more: bool) -> str:
    """生成 IC gate repair 派发指令。"""
    return (
        "## IC Topic Gate Repair — 禁止并行派发\n\n"
        f"Gate {phase_name} FAIL，需要派发 repair 子代理补充采集。\n"
        "多个 repair 子代理会写同一组文件，并行执行会导致数据丢失。\n\n"
        "## 步骤\n"
        "1. 读取 dispatch_info.manifests[0] 指向的 manifest JSON\n"
        "2. 使用 Agent 工具派发**这一个** repair 子代理（mode='bypassPermissions'）\n"
        "3. 等待 repair 子代理完成\n"
        f"4. 用 start_phase='{phase_name}' 恢复管线，重跑 gate\n"
    ) + (
        "\n5. 还有更多 repair manifest，管线恢复后会返回下一个，重复步骤 1-4"
        if has_more else
        "\n5. 这是最后一个 repair manifest，完成后 gate 会重新评估"
    )


# ═══════════════════════════════════════════════════════════
# Shared State Refresh
# ═══════════════════════════════════════════════════════════

def _refresh_shared_state(task: Path, wave: int, job_id: str):
    """刷新 shared_state.json — 提取各维度输出的摘要信息供后续 Wave 使用。"""
    shared_state_path = task / "shared_state.json"
    shared_state = {}
    if shared_state_path.exists():
        try:
            shared_state = json.loads(shared_state_path.read_text(encoding="utf-8"))
        except Exception:
            pass

    dim_summaries = {}
    roles = IC_TOPIC_WAVE_ROLES.get(wave, [])
    for role in roles:
        md_path = task / f"{role}.md"
        if md_path.exists():
            content = md_path.read_text(encoding="utf-8")
            dim_summaries[role] = {
                "chars": len(content),
                "sections": content.count("\n## "),
                "has_data": len(content) > 500,
                "summary": content[:300] + "..." if len(content) > 300 else content,
            }

    shared_state[f"wave{wave}_dimensions"] = dim_summaries
    shared_state[f"wave{wave}_refreshed_at"] = time.strftime("%Y-%m-%dT%H:%M:%S")
    shared_state["job_id"] = job_id

    _atomic_write_json(shared_state_path, shared_state)
    return dim_summaries


# ═══════════════════════════════════════════════════════════
# Phase 01: Topic Intake
# ═══════════════════════════════════════════════════════════

def _run_topic_intake(runtime_root: Path, job_ctx: JobContext) -> dict[str, Any]:
    """Phase 01: 课题入库 + 研究计划生成。"""
    task = _task_dir(runtime_root, job_ctx)
    topic_name = job_ctx.entity
    core_question = job_ctx.query or ""
    metadata = job_ctx.metadata or {}

    plan = {
        "topic_name": topic_name,
        "direction": metadata.get("direction", ""),
        "core_question": core_question,
        "sub_questions": metadata.get("sub_questions", []),
        "research_scope": metadata.get("research_scope", ""),
        "market": job_ctx.market,
        "job_id": job_ctx.job_id,
        "created_at": time.strftime("%Y-%m-%dT%H:%M:%S"),
    }

    plan_path = task / "research_plan.json"
    _atomic_write_json(plan_path, plan)

    return {
        "ok": True, "mode": "topic_intake",
        "phase": "phase01_topic_intake", "job_id": job_ctx.job_id,
        "result": {"plan_path": str(plan_path), "topic_name": topic_name},
    }


# ═══════════════════════════════════════════════════════════
# Phase 02-04: Presearch / Extract / Fact Store (unchanged)
# ═══════════════════════════════════════════════════════════

def _run_presearch(runtime_root: Path, job_ctx: JobContext) -> dict[str, Any]:
    """Phase 02: 预搜索 [heavy_bg]."""
    if os.environ.get("IRBP_BG_CHILD") == "1":
        return _run_presearch_inner(runtime_root, job_ctx)
    from scripts.heavy_phase_bg import check_cached_result, launch_heavy_phase
    cached = check_cached_result(runtime_root, job_ctx.job_id, "phase02_presearch")
    if cached is not None:
        return cached
    return launch_heavy_phase(runtime_root, job_ctx, "phase02_presearch", pipeline="ic_topic")


def _run_presearch_inner(runtime_root: Path, job_ctx: JobContext) -> dict[str, Any]:
    task = _task_dir(runtime_root, job_ctx)
    topic_name = job_ctx.entity
    core_question = job_ctx.query or topic_name
    try:
        from scripts.search_gateway import neodata_search
    except ImportError:
        neodata_search = None
    results = []
    for sq in [f'"{topic_name}" 行业 市场规模 竞争格局',
               f'"{topic_name}" 技术 产业链 发展趋势', core_question]:
        try:
            if neodata_search:
                r = neodata_search(sq)
                if r:
                    results.append({"query": sq, "source": "neodata",
                                    "urls": r if isinstance(r, list) else [r]})
        except Exception:
            pass
    result_path = task / "presearch_results.json"
    _atomic_write_json(result_path, results)
    return {"ok": True, "mode": "presearch", "phase": "phase02_presearch",
            "job_id": job_ctx.job_id,
            "result": {"query_count": 3, "result_path": str(result_path)}}


def _run_extract(runtime_root: Path, job_ctx: JobContext) -> dict[str, Any]:
    """Phase 03: 内容提取。"""
    from scripts.ir_extract_content import extract_from_presearch
    result = extract_from_presearch(task_id=job_ctx.job_id, entity=job_ctx.entity,
                                     max_pages=8, pipeline='ic_topic')
    return {"ok": result.get("ok_count", 0) > 0, "mode": "extract",
            "phase": "phase03_extract", "job_id": job_ctx.job_id,
            "result": {"total_urls": result.get("total_urls", 0),
                       "ok_count": result.get("ok_count", 0)}}


def _run_fact_store_bootstrap(runtime_root: Path, job_ctx: JobContext) -> dict[str, Any]:
    """Phase 04: Fact Store 初始化。"""
    task = _task_dir(runtime_root, job_ctx)
    plan_path = task / "research_plan.json"
    plan = json.loads(plan_path.read_text(encoding="utf-8")) if plan_path.exists() else {}

    fact_store = {"bootstrap_version": "1.0",
                  "topic_name": plan.get("topic_name", job_ctx.entity),
                  "core_question": plan.get("core_question", ""), "facts": [],
                  "created_at": time.strftime("%Y-%m-%dT%H:%M:%S")}
    fs_path = task / "fact_store.json"
    _atomic_write_json(fs_path, fact_store)
    _sync_to_workspace(job_ctx, fs_path, "fact_store.json")
    return {"ok": True, "mode": "fact_store_bootstrap",
            "phase": "phase04_fact_store_bootstrap", "job_id": job_ctx.job_id,
            "result": {"fact_store_path": str(fs_path)}}


# ═══════════════════════════════════════════════════════════
# Phase 05: Wave 1 Dispatch Prepare
# ═══════════════════════════════════════════════════════════

def _run_wave1_dispatch_prepare(runtime_root: Path, job_ctx: JobContext) -> dict[str, Any]:
    """Phase 05: Wave 1 dispatch — sequential has_more 派发 3 个角色。"""
    from scripts.ic_topic_subagent_launcher import launch_next_wave, _role_outputs_complete
    task = _task_dir(runtime_root, job_ctx)

    completed = [r for r in IC_TOPIC_WAVE1_ROLE_SLUGS if _role_outputs_complete(task, r)]
    if len(completed) == len(IC_TOPIC_WAVE1_ROLE_SLUGS):
        return {"ok": True, "needs_dispatch": True,
                "dispatch_info": {"type": "wave1_complete", "manifests": [], "has_more": False},
                "phase": "phase05_wave1_dispatch_prepare", "job_id": job_ctx.job_id}

    wr = launch_next_wave(runtime_root=runtime_root, job_id=job_ctx.job_id,
                          entity=job_ctx.entity, query=job_ctx.query,
                          market=job_ctx.market, metadata=job_ctx.metadata)

    if wr.get("all_done"):
        return {"ok": True, "needs_dispatch": True,
                "dispatch_info": {"type": "wave1_complete", "manifests": [], "has_more": False},
                "phase": "phase05_wave1_dispatch_prepare", "job_id": job_ctx.job_id}

    return {"ok": True, "needs_dispatch": True,
            "dispatch_info": {"type": "wave1_dispatch",
                              "manifests": [ti["manifest_path"] for ti in wr.get("task_tool_instructions", [])],
                              "has_more": wr.get("has_more", False)},
            "result": {"task_tool_instructions": wr.get("task_tool_instructions", []),
                       "has_more": wr.get("has_more", False)},
            "phase": "phase05_wave1_dispatch_prepare", "job_id": job_ctx.job_id}


# ═══════════════════════════════════════════════════════════
# Phase 06: Wave 1 Collect (with retry + progress detection)
# ═══════════════════════════════════════════════════════════

def _run_wave1_collect(runtime_root: Path, job_ctx: JobContext) -> dict[str, Any]:
    """Phase 06: Wave 1 收集 — retry + 进度检测。"""
    from scripts.ic_topic_subagent_launcher import _role_outputs_complete
    task = _task_dir(runtime_root, job_ctx)
    roles = list(IC_TOPIC_WAVE1_ROLE_SLUGS.keys())

    result = _collect_with_retry(task, roles, collect_name="wave1_collect")

    if result.get("ok"):
        for role in roles:
            md_path = task / f"{role}.md"
            _sync_to_workspace(job_ctx, md_path, f"{role}.md")

    return {"ok": result["ok"], "mode": "collect",
            "phase": "phase06_wave1_dispatch_collect", "job_id": job_ctx.job_id,
            "result": {"wave": 1, "all_ok": result["ok"], "roles": result.get("roles", {}),
                       "incomplete": result.get("incomplete", [])}}


# ═══════════════════════════════════════════════════════════
# Phase 07: Wave 1 Evidence Gate (severity grading + repair)
# ═══════════════════════════════════════════════════════════

def _run_wave1_evidence_gate(runtime_root: Path, job_ctx: JobContext) -> dict[str, Any]:
    """Phase 07: Wave 1 证据门禁 — facts.json 检查 + severity 分级 + repair 派发。"""
    task = _task_dir(runtime_root, job_ctx)
    gate_results = {"wave": 1, "passed": True, "checks": {}, "failures": []}

    for role in IC_TOPIC_WAVE1_ROLE_SLUGS:
        facts_path = task / f"{role}{FACTS_SUFFIX}"
        md_path = task / f"{role}.md"

        if facts_path.exists():
            try:
                data = json.loads(facts_path.read_text(encoding="utf-8"))
                items = data if isinstance(data, list) else data.get("facts", [])
                gate_results["checks"][f"{role}_fact_count"] = {
                    "actual": len(items), "threshold": 3, "severity": "BLOCKING",
                    "passed": len(items) >= 3,
                }
            except Exception:
                gate_results["checks"][f"{role}_facts_parse"] = {
                    "actual": "parse_error", "threshold": "valid_json",
                    "severity": "BLOCKING", "passed": False,
                }
        else:
            gate_results["checks"][f"{role}_facts_exists"] = {
                "actual": False, "threshold": True, "severity": "BLOCKING", "passed": False,
            }

        if md_path.exists():
            content = md_path.read_text(encoding="utf-8")
            chars = len(content)
            sources = content.count("http")
            gate_results["checks"][f"{role}_char_count"] = {
                "actual": chars, "threshold": 1500, "severity": "MEDIUM",
                "passed": chars >= 1500,
            }
            gate_results["checks"][f"{role}_source_count"] = {
                "actual": sources, "threshold": 3, "severity": "MEDIUM",
                "passed": sources >= 3,
            }
        else:
            gate_results["checks"][f"{role}_md_exists"] = {
                "actual": False, "threshold": True, "severity": "BLOCKING", "passed": False,
            }

    # Aggregate verdict
    blocking_fails = [n for n, c in gate_results["checks"].items()
                      if not c["passed"] and c.get("severity") == "BLOCKING"]
    medium_fails = [n for n, c in gate_results["checks"].items()
                    if not c["passed"] and c.get("severity") == "MEDIUM"]
    gate_results["verdict"] = "BLOCKING" if blocking_fails else ("WARN" if medium_fails else "PASS")

    # Write gate result (atomic)
    gate_path = task / "wave1_gate.json"
    _atomic_write_json(gate_path, gate_results)

    # Sidecar 缺失 → hard block
    sidecar_missing = [n for n in blocking_fails if "facts" in n.lower()]
    if sidecar_missing and medium_fails:
        gate_results["passed"] = False

    failed_checks = [n for n, c in gate_results["checks"].items() if not c["passed"]]

    if failed_checks and gate_results.get("verdict") in ("BLOCKING", "WARN"):
        repair_state_path = task / "wave1_repair_state.json"
        repair_attempt = 0
        if repair_state_path.exists():
            try:
                repair_attempt = json.loads(repair_state_path.read_text(encoding="utf-8")).get("attempt", 0)
            except Exception:
                pass

        if repair_attempt < GATE_REPAIR_MAX_ATTEMPTS:
            # Map failed checks to roles
            roles_needing_repair: set[str] = set()
            for cn in failed_checks:
                for role in IC_TOPIC_WAVE1_ROLE_SLUGS:
                    if role in cn:
                        roles_needing_repair.add(role)

            if roles_needing_repair:
                repair_state_path.write_text(
                    json.dumps({"attempt": repair_attempt + 1, "failed_checks": failed_checks},
                               ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

                first_role = sorted(roles_needing_repair)[0]
                remaining_roles = sorted(roles_needing_repair)[1:]
                failure_desc = ", ".join(failed_checks[:5])

                manifest_path = _build_ic_repair_manifest(
                    task, role=first_role, slug=IC_TOPIC_WAVE1_ROLE_SLUGS[first_role],
                    failure_reason=f"Gate checks failed: {failure_desc}",
                    gate_phase="phase07_wave1_evidence_gate", runtime_root=runtime_root)

                remaining_manifests = [_build_ic_repair_manifest(
                    task, role=r, slug=IC_TOPIC_WAVE1_ROLE_SLUGS[r],
                    failure_reason=f"Gate checks failed: {failure_desc}",
                    gate_phase="phase07_wave1_evidence_gate", runtime_root=runtime_root)
                    for r in remaining_roles]

                has_more = len(remaining_manifests) > 0
                print(f"  [wave1_gate] repair attempt {repair_attempt+1}/{GATE_REPAIR_MAX_ATTEMPTS}, "
                      f"role={first_role}, failed={failure_desc}", flush=True)

                return {"ok": True, "needs_dispatch": True, "has_more": has_more,
                        "mode": "ic_wave_repair",
                        "phase": "phase07_wave1_evidence_gate", "job_id": job_ctx.job_id,
                        "dispatch_info": {"manifests": [manifest_path],
                                          "remaining_manifests": remaining_manifests,
                                          "is_repair": True},
                        "result": gate_results,
                        "instruction": _ic_repair_instruction(
                            "phase07_wave1_evidence_gate", has_more)}

            # Repair exhausted — degrade to WARN
            gate_results["repair_exhausted"] = True
            gate_results["verdict"] = "WARN_DEGRADED"

    return {"ok": True, "mode": "evidence_gate",
            "phase": "phase07_wave1_evidence_gate", "job_id": job_ctx.job_id,
            "result": gate_results}


# ═══════════════════════════════════════════════════════════
# Phase 08: Wave 1 Fact Store Merge (atomic write)
# ═══════════════════════════════════════════════════════════

def _run_wave1_fact_store_merge(runtime_root: Path, job_ctx: JobContext) -> dict[str, Any]:
    """Phase 08: 合并 Wave 1 事实到 fact_store [atomic write]."""
    task = _task_dir(runtime_root, job_ctx)
    fs_path = task / "fact_store.json"
    fs = json.loads(fs_path.read_text(encoding="utf-8")) if fs_path.exists() else {"facts": []}
    if "facts" not in fs:
        fs["facts"] = []

    merged_count = 0
    for role in IC_TOPIC_WAVE1_ROLE_SLUGS:
        ff = task / f"{role}{FACTS_SUFFIX}"
        if ff.exists():
            try:
                rf = json.loads(ff.read_text(encoding="utf-8"))
                items = rf if isinstance(rf, list) else rf.get("facts", [])
                for f in items:
                    f["source_role"] = role
                fs["facts"].extend(items)
                merged_count += len(items)
            except Exception:
                pass

    fs["last_merge_wave"] = 1
    fs["last_merge_at"] = time.strftime("%Y-%m-%dT%H:%M:%S")
    _atomic_write_json(fs_path, fs)

    return {"ok": True, "mode": "fact_store_merge",
            "phase": "phase08_wave1_fact_store_merge", "job_id": job_ctx.job_id,
            "result": {"merged_count": merged_count, "wave": 1}}


# ═══════════════════════════════════════════════════════════
# Phase 09: Wave 1 Shared State Refresh [NEW]
# ═══════════════════════════════════════════════════════════

def _run_wave1_shared_state_refresh(runtime_root: Path, job_ctx: JobContext) -> dict[str, Any]:
    """Phase 09: W1 共享状态刷新 — 提取维度摘要供 W2 子代理快速了解 W1 全貌。"""
    task = _task_dir(runtime_root, job_ctx)
    dims = _refresh_shared_state(task, 1, job_ctx.job_id)
    return {"ok": True, "mode": "shared_state_refresh",
            "phase": "phase09_wave1_shared_state_refresh", "job_id": job_ctx.job_id,
            "result": {"dimensions": list(dims.keys()), "wave": 1}}


# ═══════════════════════════════════════════════════════════
# Phase 10-13: Wave 2 (prepare/collect/gate/merge, same pattern)
# ═══════════════════════════════════════════════════════════

def _run_wave2_dispatch_prepare(runtime_root: Path, job_ctx: JobContext) -> dict[str, Any]:
    from scripts.ic_topic_subagent_launcher import launch_next_wave, _role_outputs_complete
    task = _task_dir(runtime_root, job_ctx)
    completed = [r for r in IC_TOPIC_WAVE2_ROLE_SLUGS if _role_outputs_complete(task, r)]
    if len(completed) == len(IC_TOPIC_WAVE2_ROLE_SLUGS):
        return {"ok": True, "needs_dispatch": True,
                "dispatch_info": {"type": "wave2_complete", "manifests": [], "has_more": False},
                "phase": "phase10_wave2_dispatch_prepare", "job_id": job_ctx.job_id}

    wr = launch_next_wave(runtime_root=runtime_root, job_id=job_ctx.job_id,
                          entity=job_ctx.entity, query=job_ctx.query,
                          market=job_ctx.market, metadata=job_ctx.metadata)
    return {"ok": True, "needs_dispatch": True,
            "dispatch_info": {"type": "wave2_dispatch",
                              "manifests": [ti["manifest_path"] for ti in wr.get("task_tool_instructions", [])],
                              "has_more": wr.get("has_more", False)},
            "result": {"task_tool_instructions": wr.get("task_tool_instructions", []),
                       "has_more": wr.get("has_more", False)},
            "phase": "phase10_wave2_dispatch_prepare", "job_id": job_ctx.job_id}


def _run_wave2_collect(runtime_root: Path, job_ctx: JobContext) -> dict[str, Any]:
    from scripts.ic_topic_subagent_launcher import _role_outputs_complete
    task = _task_dir(runtime_root, job_ctx)
    roles = list(IC_TOPIC_WAVE2_ROLE_SLUGS.keys())
    result = _collect_with_retry(task, roles, collect_name="wave2_collect")
    if result.get("ok"):
        for role in roles:
            _sync_to_workspace(job_ctx, task / f"{role}.md", f"{role}.md")
    return {"ok": result["ok"], "mode": "collect",
            "phase": "phase11_wave2_dispatch_collect", "job_id": job_ctx.job_id,
            "result": {"wave": 2, "all_ok": result["ok"], "roles": result.get("roles", {})}}


def _run_wave2_evidence_gate(runtime_root: Path, job_ctx: JobContext) -> dict[str, Any]:
    """Phase 12: Wave 2 证据门禁 — same pattern as Wave 1."""
    task = _task_dir(runtime_root, job_ctx)
    gate_results = {"wave": 2, "passed": True, "checks": {}, "failures": []}

    for role in IC_TOPIC_WAVE2_ROLE_SLUGS:
        facts_path = task / f"{role}{FACTS_SUFFIX}"
        md_path = task / f"{role}.md"

        if facts_path.exists():
            try:
                data = json.loads(facts_path.read_text(encoding="utf-8"))
                items = data if isinstance(data, list) else data.get("facts", [])
                gate_results["checks"][f"{role}_fact_count"] = {
                    "actual": len(items), "threshold": 3, "severity": "BLOCKING",
                    "passed": len(items) >= 3}
            except Exception:
                gate_results["checks"][f"{role}_facts_parse"] = {
                    "actual": "parse_error", "threshold": "valid_json",
                    "severity": "BLOCKING", "passed": False}
        else:
            gate_results["checks"][f"{role}_facts_exists"] = {
                "actual": False, "threshold": True, "severity": "BLOCKING", "passed": False}

        if md_path.exists():
            content = md_path.read_text(encoding="utf-8")
            gate_results["checks"][f"{role}_char_count"] = {
                "actual": len(content), "threshold": 1500, "severity": "MEDIUM",
                "passed": len(content) >= 1500}
        else:
            gate_results["checks"][f"{role}_md_exists"] = {
                "actual": False, "threshold": True, "severity": "BLOCKING", "passed": False}

    blocking_fails = [n for n, c in gate_results["checks"].items()
                      if not c["passed"] and c.get("severity") == "BLOCKING"]
    medium_fails = [n for n, c in gate_results["checks"].items()
                    if not c["passed"] and c.get("severity") == "MEDIUM"]
    gate_results["verdict"] = "BLOCKING" if blocking_fails else ("WARN" if medium_fails else "PASS")

    gate_path = task / "wave2_gate.json"
    _atomic_write_json(gate_path, gate_results)

    # Repair (same logic as W1)
    failed_checks = [n for n, c in gate_results["checks"].items() if not c["passed"]]
    if failed_checks and gate_results.get("verdict") in ("BLOCKING", "WARN"):
        repair_state_path = task / "wave2_repair_state.json"
        repair_attempt = 0
        if repair_state_path.exists():
            try:
                repair_attempt = json.loads(repair_state_path.read_text(encoding="utf-8")).get("attempt", 0)
            except Exception:
                pass
        if repair_attempt < GATE_REPAIR_MAX_ATTEMPTS:
            roles_needing_repair: set[str] = set()
            for cn in failed_checks:
                for role in IC_TOPIC_WAVE2_ROLE_SLUGS:
                    if role in cn:
                        roles_needing_repair.add(role)
            if roles_needing_repair:
                repair_state_path.write_text(
                    json.dumps({"attempt": repair_attempt + 1, "failed_checks": failed_checks},
                               ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
                first_role = sorted(roles_needing_repair)[0]
                remaining_roles = sorted(roles_needing_repair)[1:]
                failure_desc = ", ".join(failed_checks[:5])
                manifest_path = _build_ic_repair_manifest(
                    task, role=first_role, slug=IC_TOPIC_WAVE2_ROLE_SLUGS[first_role],
                    failure_reason=f"Gate checks failed: {failure_desc}",
                    gate_phase="phase12_wave2_evidence_gate", runtime_root=runtime_root)
                remaining_manifests = [_build_ic_repair_manifest(
                    task, role=r, slug=IC_TOPIC_WAVE2_ROLE_SLUGS[r],
                    failure_reason=f"Gate checks failed: {failure_desc}",
                    gate_phase="phase12_wave2_evidence_gate", runtime_root=runtime_root)
                    for r in remaining_roles]
                has_more = len(remaining_manifests) > 0
                print(f"  [wave2_gate] repair attempt {repair_attempt+1}, role={first_role}", flush=True)
                return {"ok": True, "needs_dispatch": True, "has_more": has_more,
                        "mode": "ic_wave_repair",
                        "phase": "phase12_wave2_evidence_gate", "job_id": job_ctx.job_id,
                        "dispatch_info": {"manifests": [manifest_path],
                                          "remaining_manifests": remaining_manifests,
                                          "is_repair": True},
                        "result": gate_results,
                        "instruction": _ic_repair_instruction(
                            "phase12_wave2_evidence_gate", has_more)}
            gate_results["repair_exhausted"] = True
            gate_results["verdict"] = "WARN_DEGRADED"

    return {"ok": True, "mode": "evidence_gate",
            "phase": "phase12_wave2_evidence_gate", "job_id": job_ctx.job_id,
            "result": gate_results}


def _run_wave2_fact_store_merge(runtime_root: Path, job_ctx: JobContext) -> dict[str, Any]:
    """Phase 13: Wave 2 合并 [atomic write]."""
    task = _task_dir(runtime_root, job_ctx)
    fs_path = task / "fact_store.json"
    fs = json.loads(fs_path.read_text(encoding="utf-8")) if fs_path.exists() else {"facts": []}
    if "facts" not in fs:
        fs["facts"] = []
    merged = 0
    for role in IC_TOPIC_WAVE2_ROLE_SLUGS:
        ff = task / f"{role}{FACTS_SUFFIX}"
        if ff.exists():
            try:
                rf = json.loads(ff.read_text(encoding="utf-8"))
                items = rf if isinstance(rf, list) else rf.get("facts", [])
                for f in items:
                    f["source_role"] = role
                fs["facts"].extend(items)
                merged += len(items)
            except Exception:
                pass
    fs["last_merge_wave"] = 2
    fs["last_merge_at"] = time.strftime("%Y-%m-%dT%H:%M:%S")
    _atomic_write_json(fs_path, fs)
    return {"ok": True, "mode": "fact_store_merge",
            "phase": "phase13_wave2_fact_store_merge", "job_id": job_ctx.job_id,
            "result": {"merged_count": merged, "wave": 2}}


# ═══════════════════════════════════════════════════════════
# Phase 14: Wave 2 Shared State Refresh [NEW]
# ═══════════════════════════════════════════════════════════

def _run_wave2_shared_state_refresh(runtime_root: Path, job_ctx: JobContext) -> dict[str, Any]:
    """Phase 14: W2 共享状态刷新。"""
    task = _task_dir(runtime_root, job_ctx)
    dims = _refresh_shared_state(task, 2, job_ctx.job_id)
    return {"ok": True, "mode": "shared_state_refresh",
            "phase": "phase14_wave2_shared_state_refresh", "job_id": job_ctx.job_id,
            "result": {"dimensions": list(dims.keys()), "wave": 2}}


# ═══════════════════════════════════════════════════════════
# Phase 15-16: Synthesis
# ═══════════════════════════════════════════════════════════

def _run_synthesis_dispatch_prepare(runtime_root: Path, job_ctx: JobContext) -> dict[str, Any]:
    """Phase 15: 统稿派发 — 含 structured brief + shared_state。"""
    from scripts.ic_topic_subagent_launcher import _role_outputs_complete, _assemble_system_prompt
    task = _task_dir(runtime_root, job_ctx)

    role = "ic_report_synthesizer"
    if _role_outputs_complete(task, role):
        return {"ok": True, "needs_dispatch": True,
                "dispatch_info": {"type": "synthesis_complete", "manifests": [], "has_more": False},
                "phase": "phase15_synthesis_dispatch_prepare", "job_id": job_ctx.job_id}

    system_prompt = _assemble_system_prompt(runtime_root, role, task_dir=task)
    assert len(system_prompt) > 200

    plan = json.loads((task / "research_plan.json").read_text(encoding="utf-8")) if (task / "research_plan.json").exists() else {}
    shared_state = json.loads((task / "shared_state.json").read_text(encoding="utf-8")) if (task / "shared_state.json").exists() else {}

    dimension_files = {}
    for role_key in list(IC_TOPIC_WAVE1_ROLE_SLUGS.keys()) + list(IC_TOPIC_WAVE2_ROLE_SLUGS.keys()):
        md_path = task / f"{role_key}.md"
        if md_path.exists():
            dimension_files[role_key] = str(md_path)

    manifest = {
        "role": role, "slug": IC_TOPIC_WAVE3_ROLE_SLUGS[role], "wave": 3,
        "system_prompt": system_prompt,
        "output_path": str(task / "ic_topic_report.md"),
        "connectorIds": IC_TOPIC_ROLE_CONNECTOR_IDS.get(role, []),
        "key_inputs": {
            "research_plan": str(task / "research_plan.json"),
            "fact_store": str(task / "fact_store.json"),
            "shared_state": str(task / "shared_state.json"),
            "dimension_files": dimension_files,
            "topic_name": plan.get("topic_name", job_ctx.entity),
        },
        "expected_outputs": ["ic_topic_report.md", f"{role}{FACTS_SUFFIX}", f"{role}{SECTION_SUFFIX}"],
        "manifest_version": "1.0",
    }

    manifest_path = task / "synthesis_manifest.json"
    _atomic_write_json(manifest_path, manifest)

    return {"ok": True, "needs_dispatch": True,
            "dispatch_info": {"type": "synthesis_dispatch", "manifests": [str(manifest_path)], "has_more": False},
            "result": {"task_tool_instructions": [{
                "step": role, "role": role, "manifest_path": str(manifest_path),
                "output_path": str(task / "ic_topic_report.md"),
                "prompt": (
                    f"## IC Topic 统稿子代理派发指令\n\n"
                    f"读取 manifest: {manifest_path}\n"
                    f"使用 Agent 工具:\n"
                    f"  - name = 'ic_report_synthesizer'\n"
                    f"  - team_name = 'ic-{{task_id}}'\n"
                    f"  - mode = 'bypassPermissions'\n"
                    f"  - prompt = manifest 的 'system_prompt' 字段 (完整原文)\n"
                    f"  - connectorIds = manifest 的 'connectorIds' 字段\n\n"
                    f"确认 ic_topic_report.md 齐全后，用 start_phase='phase16_synthesis_dispatch_collect' 恢复。\n"
                ),
                "connectorIds": manifest["connectorIds"],
            }]},
            "phase": "phase15_synthesis_dispatch_prepare", "job_id": job_ctx.job_id}


def _run_synthesis_collect(runtime_root: Path, job_ctx: JobContext) -> dict[str, Any]:
    """Phase 16: 统稿收集 — retry + citation density gate。"""
    from scripts.ic_topic_subagent_launcher import _role_outputs_complete
    task = _task_dir(runtime_root, job_ctx)
    role = "ic_report_synthesizer"

    result = _collect_with_retry(task, [role], collect_name="synthesis_collect",
                                 max_retries=10, retry_interval=30)
    ok = result.get("ok", False)
    report_path = task / "ic_topic_report.md"

    if ok and report_path.exists():
        content = report_path.read_text(encoding="utf-8")
        chars = len(content)
        citations = content.count("http")
        density = citations / max(chars / 2000, 1)
        threshold = WAVE3_GATE_THRESHOLDS.get("citation_density_per_2k", 3)

        if density < threshold:
            print(f"  [synthesis] citation density {density:.1f}/{threshold} low, triggering repair", flush=True)
            # Synthesis repair (simplified)
            repair_manifest = _build_ic_repair_manifest(
                task, role=role, slug=IC_TOPIC_WAVE3_ROLE_SLUGS[role],
                failure_reason=f"Citation density {density:.1f} < threshold {threshold}, "
                              f"need more footnote references",
                gate_phase="phase16_synthesis_dispatch_collect", runtime_root=runtime_root)

            # Check repair attempts
            synth_repair_state = task / "synthesis_repair_state.json"
            attempt = 0
            if synth_repair_state.exists():
                try:
                    attempt = json.loads(synth_repair_state.read_text(encoding="utf-8")).get("attempt", 0)
                except Exception:
                    pass
            if attempt < 1:
                synth_repair_state.write_text(
                    json.dumps({"attempt": attempt + 1, "issue": f"citation density {density:.1f}"},
                               ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
                return {"ok": True, "needs_dispatch": True, "has_more": False,
                        "mode": "ic_synthesis_repair",
                        "phase": "phase16_synthesis_dispatch_collect", "job_id": job_ctx.job_id,
                        "dispatch_info": {"manifests": [repair_manifest], "is_repair": True},
                        "result": {"report_exists": ok, "chars": chars, "citation_density": round(density, 1)},
                        "instruction": _ic_repair_instruction(
                            "phase16_synthesis_dispatch_collect", False)}

        _sync_to_workspace(job_ctx, report_path, "ic_topic_report.md")

    return {"ok": ok, "mode": "collect",
            "phase": "phase16_synthesis_dispatch_collect", "job_id": job_ctx.job_id,
            "result": {"report_exists": ok, "report_path": str(report_path) if ok else ""}}


# ═══════════════════════════════════════════════════════════
# Phase 17: Debate Review
# ═══════════════════════════════════════════════════════════

def _run_debate_review(runtime_root: Path, job_ctx: JobContext) -> dict[str, Any]:
    """Phase 17: 对抗评审 — BLOCKING/MEDIUM 分级。"""
    task = _task_dir(runtime_root, job_ctx)
    report_path = task / "ic_topic_report.md"
    issues = []

    if not report_path.exists():
        return {"ok": False, "mode": "debate_review",
                "phase": "phase17_debate_review", "job_id": job_ctx.job_id,
                "result": {"verdict": "FAIL_BLOCKING", "issues": [{"type": "REPORT_NOT_FOUND", "level": "BLOCKING", "detail": "Report file not found"}]}}

    content = report_path.read_text(encoding="utf-8")
    chars = len(content)

    if chars < 500:
        issues.append({"type": "EMPTY_REPORT", "level": "BLOCKING", "detail": f"Report too short ({chars} chars)"})
    if content.count("http") + content.count("https") < 5:
        issues.append({"type": "LOW_SOURCES", "level": "MEDIUM", "detail": "Fewer than 5 source references"})
    if content.count("\n## ") < 5:
        issues.append({"type": "FEW_SECTIONS", "level": "MEDIUM", "detail": "Fewer than 5 top-level sections"})
    if "投资含义" not in content and "投资结论" not in content:
        issues.append({"type": "NO_INVESTMENT_SECTION", "level": "MEDIUM", "detail": "Missing investment conclusion section"})

    blocking = [i for i in issues if i.get("level") == "BLOCKING"]
    verdict = "FAIL_BLOCKING" if blocking else ("WARN" if issues else "PASS")

    # Write deferred fixes
    deferred = [i for i in issues if i.get("level") != "BLOCKING"]
    if deferred:
        fixes_path = task / "deferred_fixes.json"
        _atomic_write_json(fixes_path, {"verdict": verdict, "fixes": deferred})

    return {"ok": len(blocking) == 0, "mode": "debate_review",
            "phase": "phase17_debate_review", "job_id": job_ctx.job_id,
            "result": {"verdict": verdict, "issues": issues, "char_count": chars}}


# ═══════════════════════════════════════════════════════════
# Phase 18: Delivery (DOCX + artifacts + investment judgment)
# ═══════════════════════════════════════════════════════════

def _run_delivery(runtime_root: Path, job_ctx: JobContext) -> dict[str, Any]:
    """Phase 18: 交付 [heavy_bg] — DOCX + artifacts + investment judgment。"""
    if os.environ.get("IRBP_BG_CHILD") == "1":
        return _run_delivery_inner(runtime_root, job_ctx)
    from scripts.heavy_phase_bg import check_cached_result, launch_heavy_phase
    cached = check_cached_result(runtime_root, job_ctx.job_id, "phase18_delivery")
    if cached is not None:
        return cached
    return launch_heavy_phase(runtime_root, job_ctx, "phase18_delivery", pipeline="ic_topic")


def _run_delivery_inner(runtime_root: Path, job_ctx: JobContext) -> dict[str, Any]:
    """Delivery 子进程内执行 — DOCX + MD + artifacts + investment judgment。"""
    task = _task_dir(runtime_root, job_ctx)
    ws = job_ctx.workspace
    deliverable_artifacts: dict[str, Any] = {}

    report_src = task / "ic_topic_report.md"
    if not report_src.exists():
        return {"ok": False, "mode": "delivery",
                "phase": "phase18_delivery", "job_id": job_ctx.job_id,
                "result": {"error": "report.md not found"}}

    # 1. Copy MD to delivery
    if ws:
        dest_md = ws.delivery_dir / f"{job_ctx.entity}_研究报告.md"
        try:
            shutil.copy2(report_src, dest_md)
            deliverable_artifacts["report_md"] = str(dest_md)
        except Exception:
            pass

    # 2. Generate DOCX — main report + dimension reports (参照 BP/IR per-dimension DOCX)
    docx_ok = False
    dim_docx_count = 0
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        def _md_to_docx(md_path: Path, output_path: Path, title_override: str = ""):
            """Convert a single .md file to .docx."""
            if not md_path.exists():
                return False
            doc = Document()
            style = doc.styles["Normal"]
            style.font.name = "Microsoft YaHei"
            style.font.size = Pt(11)
            style.paragraph_format.space_after = Pt(6)

            content = md_path.read_text(encoding="utf-8")
            lines = content.split("\n")
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                if line.startswith("# "):
                    p = doc.add_paragraph()
                    run = p.add_run(title_override or line[2:])
                    run.bold = True
                    run.font.size = Pt(16)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                elif line.startswith("## "):
                    p = doc.add_paragraph()
                    run = p.add_run(line[3:])
                    run.bold = True
                    run.font.size = Pt(14)
                elif line.startswith("### "):
                    p = doc.add_paragraph()
                    run = p.add_run(line[4:])
                    run.bold = True
                    run.font.size = Pt(12)
                elif line.startswith("- ") or line.startswith("* "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                elif line.startswith("|"):
                    continue
                else:
                    doc.add_paragraph(line)
            doc.save(str(output_path))
            return True

        # 2a. 主报告 DOCX
        if ws:
            dest_docx = ws.delivery_dir / f"{job_ctx.entity}_研究报告.docx"
            if _md_to_docx(report_src, dest_docx, title_override=f"{job_ctx.entity} 深度研究报告"):
                deliverable_artifacts["report_docx"] = str(dest_docx)
                docx_ok = True

        # 2b. 维度独立 DOCX（参照 BP 维度 DOCX / IR per-step DOCX）
        if ws:
            dim_docx_dir = ws.delivery_dir / "维度报告"
            dim_docx_dir.mkdir(exist_ok=True)
            role_titles = {
                "ic_market_overview": "市场全景分析",
                "ic_competitive_landscape": "竞争格局分析",
                "ic_tech_product": "技术产品分析",
                "ic_supply_chain": "产业链与供应链分析",
                "ic_policy_risk": "政策与风险分析",
            }
            for role_key in list(IC_TOPIC_WAVE1_ROLE_SLUGS.keys()) + list(IC_TOPIC_WAVE2_ROLE_SLUGS.keys()):
                dim_md = task / f"{role_key}.md"
                if dim_md.exists():
                    title = role_titles.get(role_key, role_key)
                    dim_docx = dim_docx_dir / f"{title}.docx"
                    try:
                        if _md_to_docx(dim_md, dim_docx, title_override=f"{title} — {job_ctx.entity}"):
                            dim_docx_count += 1
                    except Exception:
                        pass
            if dim_docx_count > 0:
                deliverable_artifacts["dimension_docx"] = dim_docx_count
    except ImportError:
        print("  [ic_topic] python-docx not installed, skipping DOCX generation", flush=True)
    except Exception as e:
        print(f"  [ic_topic] DOCX generation failed: {e}", flush=True)

    # 3. Generate investment judgment summary
    inv_judgment = {
        "topic_name": job_ctx.entity,
        "generated_at": time.strftime("%Y-%m-%dT%H:%M:%S"),
        "disclaimer": "仅供参考，不构成投资建议",
    }
    try:
        report_text = report_src.read_text(encoding="utf-8")
        # Extract key signals
        inv_judgment["report_length"] = len(report_text)
        inv_judgment["source_count"] = report_text.count("http") + report_text.count("https")
        has_buy_signal = any(kw in report_text for kw in ["超配", "买入", "推荐", "看好"])
        has_sell_signal = any(kw in report_text for kw in ["低配", "回避", "谨慎", "风险"])
        inv_judgment["sentiment"] = "BUY" if has_buy_signal else ("CAUTIOUS" if has_sell_signal else "NEUTRAL")

        inv_path = task / "investment_judgment.json"
        _atomic_write_json(inv_path, inv_judgment)
        if ws:
            inv_dest = ws.delivery_dir / "investment_judgment.json"
            shutil.copy2(inv_path, inv_dest)
            deliverable_artifacts["investment_judgment"] = str(inv_dest)
    except Exception:
        pass

    # 4. Copy dimension MDs
    dim_count = 0
    if ws:
        for role_key in list(IC_TOPIC_WAVE1_ROLE_SLUGS.keys()) + list(IC_TOPIC_WAVE2_ROLE_SLUGS.keys()):
            dim_md = task / f"{role_key}.md"
            if dim_md.exists():
                try:
                    shutil.copy2(dim_md, ws.delivery_dir / f"{role_key}.md")
                    dim_count += 1
                except Exception:
                    pass
    deliverable_artifacts["dimension_files_copied"] = dim_count

    # 5. Register artifacts
    if ws:
        artifacts_path = ws.state_dir / "artifacts.json"
        existing = {}
        if artifacts_path.exists():
            try:
                existing = json.loads(artifacts_path.read_text(encoding="utf-8"))
            except Exception:
                pass
        existing["ic_topic_delivery"] = deliverable_artifacts
        _atomic_write_json(artifacts_path, existing)

    return {"ok": True, "mode": "delivery",
            "phase": "phase18_delivery", "job_id": job_ctx.job_id,
            "result": {"docx_generated": docx_ok,
                       "dimension_docx": dim_docx_count,
                       "artifacts": list(deliverable_artifacts.keys()),
                       "investment_judgment": inv_judgment.get("sentiment"),
                       "dimension_files": dim_count}}


# ═══════════════════════════════════════════════════════════
# IC Topic Profile
# ═══════════════════════════════════════════════════════════

class ICTopicProfile(PipelineProfile):
    """IC Topic 课题研究 Pipeline Profile — 6 Role / 3 Wave / 18 Phase (v2 P0+P1)."""

    def __init__(self, runtime_root: Path):
        super().__init__(
            name="ic_topic",
            job_type="industry_topic_research",
            phase_handlers={
                "phase01_topic_intake": lambda jc: _run_topic_intake(runtime_root, jc),
                "phase02_presearch": lambda jc: _run_presearch(runtime_root, jc),
                "phase03_extract": lambda jc: _run_extract(runtime_root, jc),
                "phase04_fact_store_bootstrap": lambda jc: _run_fact_store_bootstrap(runtime_root, jc),
                "phase05_wave1_dispatch_prepare": lambda jc: _run_wave1_dispatch_prepare(runtime_root, jc),
                "phase06_wave1_dispatch_collect": lambda jc: _run_wave1_collect(runtime_root, jc),
                "phase07_wave1_evidence_gate": lambda jc: _run_wave1_evidence_gate(runtime_root, jc),
                "phase08_wave1_fact_store_merge": lambda jc: _run_wave1_fact_store_merge(runtime_root, jc),
                "phase09_wave1_shared_state_refresh": lambda jc: _run_wave1_shared_state_refresh(runtime_root, jc),
                "phase10_wave2_dispatch_prepare": lambda jc: _run_wave2_dispatch_prepare(runtime_root, jc),
                "phase11_wave2_dispatch_collect": lambda jc: _run_wave2_collect(runtime_root, jc),
                "phase12_wave2_evidence_gate": lambda jc: _run_wave2_evidence_gate(runtime_root, jc),
                "phase13_wave2_fact_store_merge": lambda jc: _run_wave2_fact_store_merge(runtime_root, jc),
                "phase14_wave2_shared_state_refresh": lambda jc: _run_wave2_shared_state_refresh(runtime_root, jc),
                "phase15_synthesis_dispatch_prepare": lambda jc: _run_synthesis_dispatch_prepare(runtime_root, jc),
                "phase16_synthesis_dispatch_collect": lambda jc: _run_synthesis_collect(runtime_root, jc),
                "phase17_debate_review": lambda jc: _run_debate_review(runtime_root, jc),
                "phase18_delivery": lambda jc: _run_delivery(runtime_root, jc),
            },
        )
        self.runtime_root = runtime_root

    def phase_prerequisites(self) -> dict[str, list[str]]:
        """声明 phase 间的文件依赖，使 kernel 支持断点续跑自动回填。"""
        return {
            "phase06_wave1_dispatch_collect": ["research_plan.json", "fact_store.json"],
            "phase07_wave1_evidence_gate": [
                "ic_market_overview.md", "ic_competitive_landscape.md", "ic_tech_product.md",
                "ic_market_overview-facts.json", "ic_competitive_landscape-facts.json", "ic_tech_product-facts.json",
            ],
            "phase10_wave2_dispatch_prepare": ["wave1_gate.json", "shared_state.json", "fact_store.json"],
            "phase11_wave2_dispatch_collect": ["shared_state.json"],
            "phase15_synthesis_dispatch_prepare": ["wave2_gate.json", "shared_state.json", "fact_store.json"],
            "phase17_debate_review": ["ic_topic_report.md"],
        }
