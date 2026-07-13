#!/usr/bin/env python3
"""
BP 来源与参考 Word 文档生成器

从 bp_synthesis.md 末尾的"来源与参考"章节解析脚注，
生成独立的 Word 文档，方便投资人查阅所有来源链接。

来源按类型分组：天眼查 / NeoData / WebSearch / WebFetch / 其他。
每条来源带可点击超链接。
"""
from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from typing import Any


# ── 来源分类规则 ─────────────────────────────────────

_SOURCE_CATEGORIES = {
    "天眼查": ["天眼查", "tianyancha", "TYC"],
    "NeoData / 研报": ["NeoData", "neodata", "研报", "券商"],
    "WebSearch / 新闻": ["WebSearch", "web_search", "36氪", "氪", "虎嗅", "界面", "财新", "人民网",
                        "新华社", "新浪", "腾讯", "百度", "搜狐", "澎湃", "第一财经",
                        "Yole", "Gartner", "IDC", "Counterpoint"],
    "WebFetch / 官网": ["WebFetch", "web_fetch", "官网", "official"],
}


def _classify_source(name: str, url: str) -> str:
    """根据来源名称和 URL 判断分类。"""
    text = f"{name} {url}".lower()
    for category, keywords in _SOURCE_CATEGORIES.items():
        for kw in keywords:
            if kw.lower() in text:
                return category
    # URL-based fallback
    if "tianyancha" in url:
        return "天眼查"
    if any(domain in url for domain in ["eastmoney", "xueqiu", "wind"]):
        return "NeoData / 研报"
    return "其他来源"


# ── 解析 synthesis.md 中的来源 ────────────────────────

def _parse_footnotes(synthesis_md: str) -> list[dict[str, str]]:
    """从 synthesis.md 中解析 [^N]: 来源名 — URL (日期) 格式的脚注。

    Returns:
        [{"num": "1", "name": "天眼查", "url": "https://...", "date": "2026-06-10"}, ...]
    """
    footnotes: list[dict[str, str]] = []
    seen_ids: set[str] = set()

    for line in synthesis_md.split("\n"):
        stripped = line.strip()
        # Match [^N]: content
        m = re.match(r"^\[\^(\d+)\]:\s*(.+)$", stripped)
        if not m:
            continue
        fn_id = m.group(1)
        if fn_id in seen_ids:
            continue
        seen_ids.add(fn_id)

        content = m.group(2).strip()

        # Extract URL
        url_match = re.search(r"(https?://[^\s\)\]\"'>]+)", content)
        url = url_match.group(1) if url_match else ""

        # Extract name (everything before URL or dash separator)
        name = content
        if url:
            name = content.replace(url, "").strip()
        # Clean separators
        name = re.sub(r"[\s]*[—–\-]+\s*$", "", name).strip()
        name = re.sub(r"^\s*[—–\-]+\s*", "", name).strip()

        # Extract date from trailing (YYYY) or (YYYY-MM-DD)
        date = ""
        date_match = re.search(r"\((\d{4}(?:-\d{1,2}(?:-\d{1,2})?)?)\)\s*$", name)
        if date_match:
            date = date_match.group(1)
            name = name[:date_match.start()].strip().rstrip("—–- ")

        # Skip internal file references
        if re.search(r"bp_phase\d+|bp_fact_store|bp_step0|company_verify_report|bp_ocr", name):
            continue
        if not url and "BP自述" not in name and "内部" not in name:
            continue

        footnotes.append({
            "num": fn_id,
            "name": name or "未知来源",
            "url": url,
            "date": date,
        })

    return footnotes


# ── Word 超链接 helper ────────────────────────────────

def _add_hyperlink(paragraph, url: str, text: str, *, font_name: str = "PingFang SC",
                   font_size_pt: float = 10, color_rgb=None):
    """在 paragraph 末尾追加一个可点击的 Word 超链接。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    try:
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
    except Exception:
        run = paragraph.add_run(text)
        run.font.name = font_name
        run.font.size = __import__("docx.shared", fromlist=["Pt"]).Pt(font_size_pt)
        return

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hansi"), font_name)
    rPr.append(rFonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(font_size_pt * 2)))
    rPr.append(sz)

    if color_rgb:
        color_el = OxmlElement("w:color")
        color_el.set(qn("w:val"), f"{color_rgb[0]:02X}{color_rgb[1]:02X}{color_rgb[2]:02X}")
        rPr.append(color_el)

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# ── 主函数 ────────────────────────────────────────────

def build_bp_references_docx(
    synthesis_md_path: str | Path,
    output_path: str | Path,
    entity: str = "",
) -> Path:
    """解析 synthesis.md 中的来源脚注，生成独立的来源参考 Word 文档。

    Args:
        synthesis_md_path: bp_synthesis.md 文件路径
        output_path: 输出 DOCX 路径
        entity: 公司名称（用于标题）

    Returns:
        生成的 DOCX 文件路径
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    synthesis_md_path = Path(synthesis_md_path)
    output_path = Path(output_path)

    # Read synthesis markdown
    if not synthesis_md_path.exists():
        raise FileNotFoundError(f"Synthesis file not found: {synthesis_md_path}")

    md_content = synthesis_md_path.read_text(encoding="utf-8")

    # Parse footnotes
    footnotes = _parse_footnotes(md_content)

    if not footnotes:
        print(f"  ⚠️ No footnotes found in {synthesis_md_path}, skipping references DOCX")
        return output_path

    # Classify sources
    for fn in footnotes:
        fn["category"] = _classify_source(fn["name"], fn["url"])

    # Group by category
    categories: dict[str, list[dict]] = {}
    for fn in footnotes:
        cat = fn["category"]
        categories.setdefault(cat, []).append(fn)

    # Create document
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "PingFang SC"
    font.size = Pt(10)
    # Set eastAsia font
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), "PingFang SC")

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(f"{entity} BP 尽调 — 来源与参考" if entity else "BP 尽调 — 来源与参考")
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 78, 121)

    # Subtitle
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_para.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    sub_run.font.size = Pt(9)
    sub_run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()  # spacer

    # Summary
    summary_para = doc.add_paragraph()
    summary_run = summary_para.add_run(f"共 {len(footnotes)} 条来源，涵盖 {len(categories)} 个数据来源类别")
    summary_run.font.size = Pt(10)
    summary_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()  # spacer

    # Category order
    category_order = ["天眼查", "NeoData / 研报", "WebSearch / 新闻", "WebFetch / 官网", "其他来源"]

    global_idx = 0
    for cat in category_order:
        if cat not in categories:
            continue
        items = categories[cat]

        # Category heading
        heading = doc.add_heading(f"{cat}（{len(items)} 条）", level=2)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(46, 117, 182)

        for fn in items:
            global_idx += 1
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)

            # Index number
            idx_run = p.add_run(f"[{fn['num']}] ")
            idx_run.font.size = Pt(9)
            idx_run.font.bold = True
            idx_run.font.color.rgb = RGBColor(80, 80, 80)

            # Source name
            name_run = p.add_run(fn["name"])
            name_run.font.size = Pt(10)

            # Date
            if fn["date"]:
                date_run = p.add_run(f" ({fn['date']})")
                date_run.font.size = Pt(9)
                date_run.font.color.rgb = RGBColor(128, 128, 128)

            # URL as hyperlink
            if fn["url"]:
                p.add_run(" — ")
                _add_hyperlink(p, fn["url"], fn["url"], font_size_pt=8)

    # Also handle categories not in the predefined order
    for cat, items in categories.items():
        if cat in category_order:
            continue
        heading = doc.add_heading(f"{cat}（{len(items)} 条）", level=2)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(46, 117, 182)
        for fn in items:
            global_idx += 1
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            idx_run = p.add_run(f"[{fn['num']}] ")
            idx_run.font.size = Pt(9)
            idx_run.font.bold = True
            idx_run.font.color.rgb = RGBColor(80, 80, 80)
            name_run = p.add_run(fn["name"])
            name_run.font.size = Pt(10)
            if fn["date"]:
                date_run = p.add_run(f" ({fn['date']})")
                date_run.font.size = Pt(9)
                date_run.font.color.rgb = RGBColor(128, 128, 128)
            if fn["url"]:
                p.add_run(" — ")
                _add_hyperlink(p, fn["url"], fn["url"], font_size_pt=8)

    # Footer note
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run("— 来源与参考文档结束 —")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(160, 160, 160)

    # Save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"  📄 来源参考文档: {output_path} ({len(footnotes)} 条来源)")
    return output_path


# ── CLI entry ─────────────────────────────────────────

if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="Generate BP references DOCX")
    parser.add_argument("--synthesis", required=True, help="Path to bp_synthesis.md")
    parser.add_argument("--output", required=True, help="Output DOCX path")
    parser.add_argument("--entity", default="", help="Company name for title")
    args = parser.parse_args()
    build_bp_references_docx(args.synthesis, args.output, args.entity)
