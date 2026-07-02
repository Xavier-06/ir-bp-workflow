#!/usr/bin/env python3
"""build_lit_report_docx.py — LIT 管线技术评估报告 → Word 文档。

复用 build_ir_broker_report_docx.py 的排版基础设施（字体/表格/脚注/清洗），
但章节结构适配 LIT 管线的 9 章报告 + PRISMA + 质量分布附录。

用法:
    python3 scripts/build_lit_report_docx.py TASK-20260702-001 \
        --input jobs/TASK-20260702-001/report.md \
        --output jobs/TASK-20260702-001/delivery/report.docx
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)

ROOT = Path(__file__).resolve().parent.parent

# 复用 IR 管线的排版基础设施
sys.path.insert(0, str(Path(__file__).resolve().parent))
from build_ir_broker_report_docx import (
    sanitize_text,
    read_text,
    _strip_source_section,
    _set_eastasia_font_on_style,
    _set_eastasia_font_on_run,
    _add_inline_formatted_text,
    add_table_to_doc,
    convert_markdown_to_docx,
    parse_markdown_table,
)


def build_lit_docx(
    report_md_path: Path,
    output_path: Path,
    *,
    entity: str = "",
    fact_store_path: Path | None = None,
    academic_section_path: Path | None = None,
) -> dict:
    """从 report.md 构建 LIT 技术评估报告 DOCX。

    Returns: {"success": bool, "output": str, "sources": int, "paragraphs": int}
    """
    if not report_md_path.exists():
        return {"success": False, "error": f"report.md not found: {report_md_path}"}

    memo = read_text(report_md_path)
    if not memo or len(memo) < 500:
        return {"success": False, "error": f"report.md too short ({len(memo) if memo else 0} chars)"}

    # 清洗内部信息
    memo = sanitize_text(memo)

    # 剥离来源章节（脚注定义 + 参考列表）
    memo, sources = _strip_source_section(memo)

    doc = Document()

    # ── 样式 ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(11)
    _set_eastasia_font_on_style(style, "宋体")
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.2

    for level, size_pt, sp_before in [(1, 14, 24), (2, 13, 10), (3, 12, 8)]:
        h_style = doc.styles[f"Heading {level}"]
        h_style.font.name = "Microsoft YaHei"
        h_style.font.size = Pt(size_pt)
        h_style.font.bold = True
        h_style.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        _set_eastasia_font_on_style(h_style, "宋体")
        h_style.paragraph_format.space_before = Pt(sp_before)
        h_style.paragraph_format.space_after = Pt(4)

    for list_style_name in ["List Bullet", "List Number"]:
        try:
            ls = doc.styles[list_style_name]
            ls.font.name = "Microsoft YaHei"
            ls.font.size = Pt(10.5)
            _set_eastasia_font_on_style(ls, "宋体")
            ls.paragraph_format.space_after = Pt(3)
        except KeyError:
            pass

    # ── 封面 ──
    for _ in range(4):
        doc.add_paragraph("")

    title_text = entity or "技术评估报告"
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run(f"{title_text} 技术评估报告")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    _set_eastasia_font_on_run(run, "宋体")

    doc.add_paragraph("")

    meta = doc.add_paragraph()
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = meta.add_run(f"生成日期: {datetime.now().strftime('%Y年%m月%d日')}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    _set_eastasia_font_on_run(run, "宋体")

    conf = doc.add_paragraph()
    conf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = conf.add_run("内部研究讨论稿 — 非投资建议")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)
    run.font.italic = True
    _set_eastasia_font_on_run(run, "宋体")

    doc.add_page_break()

    # ── 主体内容 ──
    convert_markdown_to_docx(memo, doc)

    # ── 来源与参考 ──
    doc.add_page_break()
    doc.add_heading("参考文献", level=1)

    if sources:
        footnote_sources = [s for s in sources if s.get("is_footnote")]
        table_sources = [s for s in sources if not s.get("is_footnote")]

        if footnote_sources:
            try:
                footnote_sources.sort(key=lambda s: int(s.get("id", "0")))
            except (ValueError, TypeError):
                pass

            doc.add_paragraph("以下为本报告引用的全部外部来源。正文中 [N] 上标标记对应此处编号。")
            doc.add_paragraph("")

            for src in footnote_sources:
                fn_id = src.get("id", "?")
                name = src.get("name", "")
                url = src.get("url", "")

                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)

                run_id = p.add_run(f"[{fn_id}] ")
                run_id.font.bold = True
                run_id.font.name = "Microsoft YaHei"
                run_id.font.size = Pt(10)
                run_id.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
                _set_eastasia_font_on_run(run_id, "宋体")

                if name:
                    run_name = p.add_run(name)
                    run_name.font.name = "Microsoft YaHei"
                    run_name.font.size = Pt(10)
                    _set_eastasia_font_on_run(run_name, "宋体")

                if url:
                    if name:
                        run_sep = p.add_run(" — ")
                        run_sep.font.name = "Microsoft YaHei"
                        run_sep.font.size = Pt(10)
                        run_sep.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                        _set_eastasia_font_on_run(run_sep, "宋体")
                    run_url = p.add_run(url)
                    run_url.font.name = "Microsoft YaHei"
                    run_url.font.size = Pt(9)
                    run_url.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
                    _set_eastasia_font_on_run(run_url, "宋体")

        if table_sources:
            source_rows = [["编号", "来源", "URL", "用途"]]
            start_idx = len(footnote_sources) + 1
            for idx, src in enumerate(table_sources, start_idx):
                source_rows.append([
                    f"[{idx}]",
                    src.get("name", ""),
                    src.get("url", ""),
                    src.get("usage", ""),
                ])
            add_table_to_doc(doc, source_rows)
    else:
        # fallback: 从正文提取 fact_id 引用
        fact_ids = sorted(set(re.findall(r"\[(READ-\d+|IND-\d+|ENT-\d+)\]", memo)))
        if fact_ids:
            doc.add_paragraph("本报告引用了以下 evidence IDs：")
            for fid in fact_ids:
                p = doc.add_paragraph()
                run = p.add_run(f"[{fid}]")
                run.font.bold = True
                run.font.name = "Microsoft YaHei"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
                _set_eastasia_font_on_run(run, "宋体")
        else:
            doc.add_paragraph("详见正文中的来源标注。")

    # ── PRISMA 附录（如有） ──
    if academic_section_path and academic_section_path.exists():
        try:
            asec = json.loads(academic_section_path.read_text(encoding="utf-8"))
            prisma = asec.get("prisma_funnel", {})
            if prisma:
                doc.add_page_break()
                doc.add_heading("附录: PRISMA 文献筛选流程", level=1)
                rows = [
                    ["阶段", "数量"],
                    ["Identification (全库命中)", str(prisma.get("identification_total", "N/A"))],
                    ["Duplicates removed", str(prisma.get("duplicates_removed", "N/A"))],
                    ["Screening excluded", str(prisma.get("screening_excluded", "N/A"))],
                    ["Included (最终纳入)", str(prisma.get("included", "N/A"))],
                ]
                add_table_to_doc(doc, rows)
        except Exception:
            pass

    # ── 免责声明 ──
    doc.add_page_break()
    doc.add_heading("免责声明", level=1)
    disclaimers = [
        "本报告由 AI 系统辅助生成，仅供内部研究讨论使用，不构成任何投资建议。",
        "报告中的数据来源于公开信息，可能存在滞后或偏差。",
        "使用者应自行核实关键数据并独立做出投资决策。",
    ]
    for d in disclaimers:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(d)
        run.font.name = "Microsoft YaHei"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        _set_eastasia_font_on_run(run, "宋体")

    # ── 保存 ──
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    return {
        "success": True,
        "output": str(output_path),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "sources_extracted": len(sources),
    }


def main():
    ap = argparse.ArgumentParser(description="LIT 管线技术评估报告 → DOCX")
    ap.add_argument("task_id", help="任务 ID (e.g. TASK-20260702-001)")
    ap.add_argument("--input", help="report.md 路径 (默认: jobs/{task_id}/report.md)")
    ap.add_argument("--output", help="输出 DOCX 路径 (默认: jobs/{task_id}/delivery/report.docx)")
    ap.add_argument("--entity", default="", help="技术名称 (封面标题)")
    ap.add_argument("--fact-store", help="fact_store.json 路径")
    ap.add_argument("--academic-section", help="academic_scout-section.json 路径 (PRISMA)")
    args = ap.parse_args()

    jobs_dir = ROOT / "runtime" / "jobs"
    task_dir = jobs_dir / args.task_id

    report_path = Path(args.input) if args.input else task_dir / "report.md"
    output_path = Path(args.output) if args.output else task_dir / "delivery" / "report.docx"

    fact_store = Path(args.fact_store) if args.fact_store else task_dir / "fact_store.json"
    acad_section = Path(args.academic_section) if args.academic_section else task_dir / "academic_scout-section.json"

    result = build_lit_docx(
        report_path,
        output_path,
        entity=args.entity,
        fact_store_path=fact_store if fact_store.exists() else None,
        academic_section_path=acad_section if acad_section.exists() else None,
    )

    print(json.dumps(result, ensure_ascii=False, indent=2))
    if not result.get("success"):
        sys.exit(1)


if __name__ == "__main__":
    main()
