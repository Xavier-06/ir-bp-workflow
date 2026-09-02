#!/usr/bin/env python3
"""
build_ic_industry_report_docx.py — 将 IC 管线 step_master_synthesis 输出转为行业深度研报 Word 文档

与 build_ir_broker_report_docx.py 对标，但章节结构适配行业研究：
  - 投资摘要 → 行业概览 → 产业链分析 → 竞争格局 → 技术趋势 → 市场规模
  - 财务基准 → 估值基准 → 资本动向 → 跨环节对比 → 投资机会 → 风险评估 → 来源附录

v1 (2026-05-29): 初始版本，复用 IR 的字体/清洗基础设施
"""
from __future__ import annotations

import argparse
import json
import re
import shutil
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

ROOT = Path(__file__).resolve().parent.parent
TASKS = ROOT / 'data' / 'tasks'
JOBS = ROOT / 'jobs'
REPORTS = ROOT / 'reports'


# ─── 东亚字体设置 ────────────────────────────

def _set_eastasia_font_on_style(style, font_name: str):
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def _set_eastasia_font_on_run(run, font_name: str):
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


# ─── 内部信息清洗 ────────────────────────────

INTERNAL_PATTERNS = [
    (r'/Users/\w+/[^\s\n]+', ''),
    (r'/home/\w+/[^\s\n]+', ''),
    (r'~/.workbuddy/[^\s\n]+', ''),
    (r'data/tasks/[^\s\n]+', ''),
    (r'TASK-\d{8}-\d{3}[-\w]*', ''),
    (r'python3?\s+scripts/[^\n]+', ''),
    (r'step\d+_\w+\.md', ''),
    (r'输出文件[：:][^\n]+', ''),
    (r'^[-*]\s*(?:Task|Entity|Accepted evidence|Rounds|Generated)[：:][^\n]*$', ''),
]

LINE_DELETE_PATTERNS = [
    r'^\s*(?:输出文件|Output file|Task ID|任务 ID)\s*[：:]',
    r'^\s*```\s*(?:bash|python|shell)',
    r'^\s*```\s*$',
    r'^\s*(?:子代理|subagent|sub-agent)',
]

_INTERNAL_SOURCE_PATTERNS = [
    r"step_\w+\.md",
    r"ic_presearch",
    r"ic_extract",
    r"data/tasks/",
]


def _is_internal_source_row(row_text: str) -> bool:
    for pat in _INTERNAL_SOURCE_PATTERNS:
        if re.search(pat, row_text, re.IGNORECASE):
            return True
    return False


def sanitize_text(text: str) -> str:
    for pat, repl in INTERNAL_PATTERNS:
        text = re.sub(pat, repl, text)
    lines = text.split('\n')
    cleaned = []
    for line in lines:
        skip = False
        for pat in LINE_DELETE_PATTERNS:
            if re.search(pat, line):
                skip = True
                break
        if not skip:
            cleaned.append(line)
    return '\n'.join(cleaned).strip()


# ─── Markdown 解析 → DOCX ────────────────────────────

FONT_NAME = 'Microsoft YaHei'
FONT_NAME_ASCII = 'Calibri'

# IC 研报配色（行业研究风格：深蓝+橙）
COLOR_TITLE = RGBColor(0x1A, 0x3C, 0x6E)       # 深蓝
COLOR_SECTION = RGBColor(0x2C, 0x5F, 0x8A)      # 中蓝
COLOR_ACCENT = RGBColor(0xE8, 0x6C, 0x00)        # 橙色
COLOR_BODY = RGBColor(0x33, 0x33, 0x33)          # 深灰
COLOR_TABLE_HEADER = RGBColor(0x1A, 0x3C, 0x6E)  # 深蓝表头


def _add_run(para, text, bold=False, italic=False, color=None, size=None):
    run = para.add_run(text)
    run.font.name = FONT_NAME_ASCII
    _set_eastasia_font_on_run(run, FONT_NAME)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return run


def _clean_heading_text(text: str) -> str:
    """去掉标题中的内部路线代号（route_xxx），用于目录显示与锚点统一"""
    text = re.sub(r'（route_[^）]*）', '', text)
    text = re.sub(r'\(route_[^)]*\)', '', text)
    return text.strip()


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT_NAME_ASCII
        _set_eastasia_font_on_run(run, FONT_NAME)
        if level == 1:
            run.font.color.rgb = COLOR_TITLE
        elif level == 2:
            run.font.color.rgb = COLOR_SECTION
    # 给标题加书签（与目录 anchor 同算法：基于去 route 后的标题），供目录超链接跳转
    try:
        anchor_text = _clean_heading_text(text)
        anchor = f'toc_{abs(hash(anchor_text)) % 100000}'
        bm_id = str(abs(hash(anchor)) % 100000)
        start = h._p.makeelement(qn('w:bookmarkStart'), {qn('w:id'): bm_id, qn('w:name'): anchor})
        end = h._p.makeelement(qn('w:bookmarkEnd'), {qn('w:id'): bm_id})
        h._p.insert(0, start)
        h._p.append(end)
    except Exception:
        pass
    return h


def _add_body_para(doc, text, bold=False):
    p = doc.add_paragraph()
    _add_run(p, text, bold=bold, color=COLOR_BODY, size=10.5)
    p.paragraph_format.space_after = Pt(4)
    return p


def _add_image(doc, img_path: str, caption: str = ''):
    """插入图片（居中，宽15.5cm）+ 可选图注（居中灰色小字）"""
    p = Path(img_path)
    if not p.exists():
        _add_body_para(doc, f"[图片缺失: {caption or img_path}]")
        return
    para = doc.add_paragraph()
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para.paragraph_format.space_before = Pt(6)
    run = para.add_run()
    run.add_picture(str(p), width=Cm(15.5))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        _add_run(cap, caption, color=RGBColor(0x66, 0x66, 0x66), size=9)
        cap.paragraph_format.space_after = Pt(8)
    return para


# ─── 编号修复（每个有序列表独立计数）────────────────────────

# python-docx 全局状态：已分配的 abstractNum/num id（跨 build 调用累积，
# python-docx 每次新建文档会重新分配 numbering，旧 id 不会冲突）
_CLONED_LIST_NUM_IDS: set = set()


def _find_next_list_num_id(num_part, taken: set) -> int:
    """在 numbering.xml 中找一个未被 num/abstractNum 占用的新 id"""
    used = set(taken)
    for child in num_part.element.iterchildren():
        if child.tag == qn('w:num'):
            nid = child.get(qn('w:numId'))
            if nid:
                used.add(int(nid))
        elif child.tag == qn('w:abstractNum'):
            aid = child.get(qn('w:abstractNumId'))
            if aid:
                used.add(int(aid))
    cand = (max(used) if used else 0) + 1
    while cand in used:
        cand += 1
    return cand


def _new_numbering_instance(p):
    """克隆 List Number 样式的编号定义 → 返回一个全新 numId（独立计数）"""
    try:
        style = p.style
        num_id_src = None
        if style is not None and style.element is not None:
            pPr = style.element.find(qn('w:pPr'))
            if pPr is not None:
                numPr = pPr.find(qn('w:numPr'))
                if numPr is not None:
                    el = numPr.find(qn('w:numId'))
                    if el is not None and el.get(qn('w:val')):
                        num_id_src = int(el.get(qn('w:val')))
        if num_id_src is None:
            return None
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        numbering_part = p.part.part_related_by(RT.NUMBERING)
    except Exception:
        return None

    # 找源 abstractNum（含多级定义）
    src_abstract = None
    for child in numbering_part.element.iterchildren():
        if child.tag == qn('w:num') and child.get(qn('w:numId')) == str(num_id_src):
            ref = child.find(qn('w:abstractNumId'))
            if ref is not None:
                target_aid = ref.get(qn('w:val'))
                for c2 in numbering_part.element.iterchildren():
                    if c2.tag == qn('w:abstractNum') and c2.get(qn('w:abstractNumId')) == target_aid:
                        src_abstract = c2
                        break
            break
    if src_abstract is None:
        return None

    taken = set(_CLONED_LIST_NUM_IDS)
    new_aid = _find_next_list_num_id(numbering_part, taken)
    taken.add(new_aid)
    new_nid = _find_next_list_num_id(numbering_part, taken)

    # 深拷贝 abstractNum，重设 id（OOXML 要求 abstractNum 在 num 之前）
    import copy
    new_abstract = copy.deepcopy(src_abstract)
    new_abstract.set(qn('w:abstractNumId'), str(new_aid))
    nsid_el = new_abstract.find(qn('w:nsid'))
    if nsid_el is not None:
        new_abstract.remove(nsid_el)
    last_abstract = None
    for child in numbering_part.element.iterchildren():
        if child.tag == qn('w:abstractNum'):
            last_abstract = child
    if last_abstract is not None:
        last_abstract.addnext(new_abstract)
    else:
        numbering_part.element.insert(0, new_abstract)

    # 新 num 引用新 abstractNum（追加到所有 num 之后）
    num_el = numbering_part.element.makeelement(qn('w:num'), {qn('w:numId'): str(new_nid)})
    ref_el = numbering_part.element.makeelement(qn('w:abstractNumId'), {qn('w:val'): str(new_aid)})
    num_el.append(ref_el)
    numbering_part.element.append(num_el)

    _CLONED_LIST_NUM_IDS.add(new_aid)
    _CLONED_LIST_NUM_IDS.add(new_nid)
    return new_nid


def _attach_numbering(p, num_id: int):
    """段落 pPr 直接挂指定 numId（level 0）"""
    pPr = p._p.get_or_add_pPr()
    numPr = pPr.find(qn('w:numPr'))
    if numPr is None:
        numPr = pPr.makeelement(qn('w:numPr'), {})
        pPr.insert(0, numPr)
    for tag in ('w:ilvl', 'w:numId'):
        el = numPr.find(qn(tag))
        if el is not None:
            numPr.remove(el)
    ilvl = numPr.makeelement(qn('w:ilvl'), {qn('w:val'): '0'})
    numid = numPr.makeelement(qn('w:numId'), {qn('w:val'): str(num_id)})
    numPr.append(ilvl)
    numPr.append(numid)


# ─── 目录（TOC 域 + 可点击书签超链接）────────────────────────

def _add_bookmark(doc, anchor: str):
    """在当前位置（段落开头）插入书签，供目录超链接跳转"""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(0)
    start = para._p.makeelement(qn('w:bookmarkStart'), {qn('w:id'): str(abs(hash(anchor)) % 100000), qn('w:name'): anchor})
    end = para._p.makeelement(qn('w:bookmarkEnd'), {qn('w:id'): str(abs(hash(anchor)) % 100000)})
    para._p.append(start)
    para._p.append(end)
    return para


def _add_toc_link(doc, text: str, anchor: str, level: int):
    """目录项：带书签超链接的段落（点击跳转到正文对应标题）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    if level == 1:
        p.paragraph_format.left_indent = Cm(0)
    else:
        p.paragraph_format.left_indent = Cm(0.8)
    # 超链接 run
    run = p.add_run(text)
    run.font.name = FONT_NAME_ASCII
    _set_eastasia_font_on_run(run, FONT_NAME)
    run.bold = (level == 1)
    run.font.size = Pt(11 if level == 1 else 10)
    run.font.color.rgb = COLOR_TITLE if level == 1 else COLOR_BODY
    # 内部超链接关系
    part = doc.part
    r_id = part.relate_to(part, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
                          is_external=False)
    # 构造 w:hyperlink
    hl = p._p.makeelement(qn('w:hyperlink'), {qn('w:anchor'): anchor})
    r = p._p.makeelement(qn('w:r'), {})
    rPr = p._p.makeelement(qn('w:rPr'), {})
    # 超链接样式（蓝色+下划线）
    rStyle = p._p.makeelement(qn('w:rStyle'), {qn('w:val'): 'Hyperlink'})
    rPr.append(rStyle)
    r.append(rPr)
    t = p._p.makeelement(qn('w:t'), {qn('xml:space'): 'preserve'})
    t.text = text
    r.append(t)
    hl.append(r)
    # 清空默认 run，用 hyperlink 替换
    for child in list(p._p):
        if child.tag != qn('w:pPr'):
            p._p.remove(child)
    p._p.append(hl)
    return p


def _build_toc_field(doc, md_text: str = ''):
    """在封面分页后插入目录页：静态目录 + 书签超链接（点击可跳转正文）"""
    # 目录标题
    toc_heading = doc.add_paragraph()
    toc_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _add_run(toc_heading, '目  录', bold=True, color=COLOR_TITLE, size=16)
    toc_heading.paragraph_format.space_after = Pt(12)

    # 从 markdown 提取 H2/H3 标题生成静态目录（含锚点）
    headings = []
    if md_text:
        for line in md_text.split('\n'):
            m2 = re.match(r'^## (.+)$', line)
            m3 = re.match(r'^### (.+)$', line)
            if m2:
                headings.append((1, m2.group(1).strip()))
            elif m3:
                txt = m3.group(1).strip()
                # 去掉内部路线代号（route_xxx），目录显示干净名称
                txt = re.sub(r'（route_[^）]*）', '', txt)
                txt = re.sub(r'\(route_[^)]*\)', '', txt)
                if len(txt) <= 40:
                    headings.append((2, txt))

    if headings:
        for level, title in headings:
            cleaned = _clean_heading_text(title)
            anchor = f'toc_{abs(hash(cleaned)) % 100000}'
            _add_toc_link(doc, cleaned, anchor, level)
    else:
        # 兜底：TOC 域
        p = doc.add_paragraph()
        r1 = p.add_run()
        fld1 = r1._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin', qn('w:dirty'): 'true'})
        r1._r.append(fld1)
        r2 = p.add_run()
        instr = r2._r.makeelement(qn('w:instrText'), {})
        instr.set(qn('xml:space'), 'preserve')
        instr.text = ' TOC \\o "1-3" \\h \\z \\u '
        r2._r.append(instr)
        r3 = p.add_run()
        fld2 = r3._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
        r3._r.append(fld2)

    doc.add_page_break()


def _add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        _add_run(p, h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=10)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Background color
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{COLOR_TABLE_HEADER.__str__()}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            if c_idx < len(headers):
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = ''
                p = cell.paragraphs[0]
                _add_run(p, str(val), color=COLOR_BODY, size=9.5)

    return table


def _parse_markdown_to_docx(doc, md_text: str):
    """将 Markdown 文本增量解析并写入 docx Document"""
    # 先提取图片行（sanitize 会删除 /Users/... 内部路径，必须在 sanitize 前暂存）
    _img_stash = {}

    def _stash_img(m):
        key = f'@@IMG{len(_img_stash)}@@'
        _img_stash[key] = m.group(0)
        return key
    md_text = re.sub(r'^!\[.*?\]\(.*?\)\s*$', _stash_img, md_text, flags=re.M)

    md_text = sanitize_text(md_text)
    lines = md_text.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]

        # 空行跳过
        if not line.strip():
            i += 1
            continue

        # 图片占位符 @@IMGn@@
        m_tok = re.match(r'^@@IMG(\d+)@@\s*$', line.strip())
        if m_tok:
            orig = _img_stash.get(f'@@IMG{m_tok.group(1)}@@', '')
            mm = re.match(r'^!\[(.*?)\]\((.*?)\)\s*$', orig)
            if mm:
                _add_image(doc, mm.group(2), mm.group(1))
            i += 1
            continue

        # 标题
        if line.startswith('#'):
            level = len(re.match(r'^#+', line).group())
            title_text = line.lstrip('#').strip()
            _add_heading(doc, title_text, level=min(level, 4))
            i += 1
            continue

        # 表格
        if '|' in line and line.strip().startswith('|'):
            # 收集连续表格行
            table_lines = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            # 解析
            headers = [c.strip() for c in table_lines[0].split('|')[1:-1] if c.strip()]
            rows = []
            for tl in table_lines[2:]:  # 跳过分隔行
                cells = [c.strip() for c in tl.split('|')[1:-1]]
                rows.append(cells)
            if headers:
                _add_table(doc, headers, rows)
            continue

        # 列表项
        if re.match(r'^\s*[-*]\s', line):
            text = re.sub(r'^\s*[-*]\s', '', line).strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # 去加粗标记
            p = doc.add_paragraph(style='List Bullet')
            _add_run(p, text, color=COLOR_BODY, size=10.5)
            i += 1
            continue

        # 有序列表（按连续块分组，每个块独立编号计数，避免全文串号）
        if re.match(r'^\s*\d+\.\s', line):
            # 收集连续的有序列表行（允许编号项之间存在空行）
            block = []
            j = i
            while j < len(lines):
                if re.match(r'^\s*\d+\.\s', lines[j]):
                    block.append(lines[j])
                    j += 1
                elif not lines[j].strip():
                    # 空行：仅当后面紧跟编号项时才视为块内分隔
                    k = j
                    while k < len(lines) and not lines[k].strip():
                        k += 1
                    if k < len(lines) and re.match(r'^\s*\d+\.\s', lines[k]):
                        j = k
                    else:
                        break
                else:
                    break
            i = j
            # 为本块创建独立编号实例
            first_p = doc.add_paragraph(style='List Number')
            first_text = re.sub(r'^\s*\d+\.\s', '', block[0]).strip()
            first_text = re.sub(r'\*\*(.*?)\*\*', r'\1', first_text)
            _add_run(first_p, first_text, color=COLOR_BODY, size=10.5)
            num_id = _new_numbering_instance(first_p)
            if num_id:
                _attach_numbering(first_p, num_id)
            for blk_line in block[1:]:
                text = re.sub(r'^\s*\d+\.\s', '', blk_line).strip()
                text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
                p = doc.add_paragraph(style='List Number')
                _add_run(p, text, color=COLOR_BODY, size=10.5)
                if num_id:
                    _attach_numbering(p, num_id)
            continue

        # 引用块
        if line.startswith('>'):
            text = line.lstrip('>').strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            p = doc.add_paragraph()
            _add_run(p, text, italic=True, color=RGBColor(0x66, 0x66, 0x66), size=10)
            p.paragraph_format.left_indent = Cm(1)
            i += 1
            continue

        # 普通段落
        text = line.strip()
        # 处理加粗标记
        if '**' in text:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    _add_run(p, part[2:-2], bold=True, color=COLOR_BODY, size=10.5)
                else:
                    _add_run(p, part, color=COLOR_BODY, size=10.5)
        else:
            _add_body_para(doc, text)

        i += 1


# ─── 主构建流程 ────────────────────────────

def build_ic_report(task_id: str, output_path: str = '') -> str:
    """将 IC 管线 step_master_synthesis 输出转为行业深度研报 DOCX"""

    # 1. 查找 master_synthesis 输出
    master_md = None
    # Workspace 模式
    ws_output = JOBS / task_id / 'outputs' / 'step_master_synthesis.md'
    if ws_output.exists() and ws_output.stat().st_size > 100:
        master_md = ws_output.read_text(encoding='utf-8')
    # Legacy 模式
    if not master_md:
        legacy = TASKS / f'{task_id}-step_master_synthesis.md'
        if legacy.exists() and legacy.stat().st_size > 100:
            master_md = legacy.read_text(encoding='utf-8')

    if not master_md:
        raise FileNotFoundError(f"step_master_synthesis 输出未找到 (task_id={task_id})")

    # 2. 读取行业元数据
    entity = ''
    scope_path = TASKS / f'{task_id}-ic_scope.json'
    if scope_path.exists():
        try:
            scope = json.loads(scope_path.read_text(encoding='utf-8'))
            entity = scope.get('industry', '')
        except Exception:
            pass

    # 3. 创建文档
    doc = Document()

    # 设置默认样式
    for style_name in ('Normal', 'Heading 1', 'Heading 2', 'Heading 3'):
        try:
            style = doc.styles[style_name]
            _set_eastasia_font_on_style(style, FONT_NAME)
            style.font.name = FONT_NAME_ASCII
        except KeyError:
            pass

    # 4. 封面页
    doc.add_paragraph()
    doc.add_paragraph()
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _add_run(title_para, f'{entity}行业深度研究报告', bold=True, color=COLOR_TITLE, size=26)

    doc.add_paragraph()
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _add_run(subtitle_para, 'Industry Deep Dive Report', italic=True, color=COLOR_SECTION, size=14)

    doc.add_paragraph()
    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _add_run(date_para, f'{datetime.now().strftime("%Y年%m月%d日")}', color=RGBColor(0x66, 0x66, 0x66), size=11)

    disclaimer_para = doc.add_paragraph()
    disclaimer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _add_run(disclaimer_para, '本报告由 AI 投研管线自动生成，仅供参考，不构成投资建议',
             italic=True, color=RGBColor(0x99, 0x99, 0x99), size=9)

    doc.add_page_break()

    # 4.5 目录页（静态目录：从 markdown 标题直接生成，任何环境下可见）
    _build_toc_field(doc, master_md)

    # 5. 正文
    _parse_markdown_to_docx(doc, master_md)

    # 6. 保存
    REPORTS.mkdir(parents=True, exist_ok=True)
    if not output_path:
        output_path = str(REPORTS / f'{task_id}_行业研报_{entity}.docx')

    doc.save(output_path)

    # 7. 复制到桌面
    try:
        desktop = Path.home() / 'Desktop'
        if desktop.exists():
            desktop_copy = desktop / Path(output_path).name
            shutil.copy2(output_path, desktop_copy)
            print(f"  📋 已复制到桌面: {desktop_copy}", flush=True)
    except Exception:
        pass

    return output_path


# ─── CLI ────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="IC 行业深度研报 DOCX 生成")
    parser.add_argument('task_id', help='任务 ID (如 TASK-20260529-001)')
    parser.add_argument('-o', '--output', default='', help='输出路径')
    args = parser.parse_args()

    try:
        out = build_ic_report(args.task_id, args.output)
        print(f"✅ 行业研报已生成: {out}")
    except Exception as e:
        print(f"❌ 生成失败: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == '__main__':
    main()
