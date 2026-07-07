#!/usr/bin/env python3
"""IR 管线 — 单 step markdown → DOCX 轻量渲染器（per-step DOCX 独立导出）。

用于将每个子代理 step 的 markdown 报告（{job_id}-{step}.md）转为独立 Word 文档，
便于增量交付与审阅。复用 build_ir_broker_report_docx.sanitize_text 的内部信息清洗规则，
保证与整篇研报 DOCX 一致的来源脱敏质量。

仅支持常见 markdown 子集：标题、段落、无序/有序列表（含一级缩进）、
简单表格、引用块、分隔线、粗体/斜体、[^N] 脚注上标。复杂场景（嵌套表格等）
按最佳努力渲染。
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

# 默认东亚字体（与 broker 构建器对齐，解决 macOS 渲染回退问题）
_EASTASIA_FONT = "Microsoft YaHei"

# 复用整篇研报构建器的脱敏清洗
try:
    from scripts.build_ir_broker_report_docx import sanitize_text  # type: ignore
except Exception:  # pragma: no cover
    sanitize_text = None  # type: ignore


def _set_eastasia(run, font: str = _EASTASIA_FONT) -> None:
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font)


def _apply_base_style(doc: Document) -> None:
    style = doc.styles['Normal']
    style.font.name = _EASTASIA_FONT
    style.font.size = Pt(10.5)
    _set_eastasia(style.font, _EASTASIA_FONT)


def _add_inline(paragraph, text: str) -> None:
    """处理行内 **粗体** / *斜体* / [^N] 脚注上标。"""
    pattern = re.compile(r'\*\*(.+?)\*\*|\[\^(\d+)\]|[*](.+?)[*]')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            _set_eastasia(r)
        if m.group(1) is not None:
            r = paragraph.add_run(m.group(1))
            r.bold = True
            _set_eastasia(r)
        elif m.group(2) is not None:
            r = paragraph.add_run(f"[{m.group(2)}]")
            r.font.superscript = True
            _set_eastasia(r)
        elif m.group(3) is not None:
            r = paragraph.add_run(m.group(3))
            r.italic = True
            _set_eastasia(r)
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        _set_eastasia(r)


def _render_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = 'Light Grid Accent 1'
    for i, row in enumerate(rows):
        for j in range(n_cols):
            cell = table.cell(i, j)
            cell.text = ""
            val = row[j].strip() if j < len(row) else ""
            # 单元格内简单去粗体标记
            val = re.sub(r'\*\*(.+?)\*\*', r'\1', val)
            run = cell.paragraphs[0].add_run(val)
            _set_eastasia(run)
            if i == 0:
                run.bold = True
    doc.add_paragraph()


def _render_markdown(doc: Document, text: str, title: str = "") -> None:
    if title:
        h = doc.add_heading(title, level=1)
        for r in h.runs:
            _set_eastasia(r)

    lines = text.split("\n")
    i = 0
    n = len(lines)
    in_code = False
    code_buffer: list[str] = []
    table_buffer: list[list[str]] = []

    def _flush_table():
        nonlocal table_buffer
        if table_buffer:
            _render_table(doc, table_buffer)
            table_buffer = []

    while i < n:
        raw = lines[i]
        stripped = raw.strip()

        # 代码块
        if stripped.startswith("```"):
            if not in_code:
                _flush_table()
                in_code = True
                code_buffer = []
            else:
                in_code = False
                code_text = "\n".join(code_buffer)
                p = doc.add_paragraph()
                p.style = doc.styles['Normal']
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            i += 1
            continue
        if in_code:
            code_buffer.append(raw)
            i += 1
            continue

        # 空行
        if not stripped:
            _flush_table()
            i += 1
            continue

        # 分隔线
        if re.match(r'^[-*_]{3,}$', stripped):
            _flush_table()
            doc.add_paragraph("_" * 20)
            i += 1
            continue

        # 表格（当前行含 | 且下一行是分隔行）
        if '|' in stripped and i + 1 < n and re.match(r'^\s*\|?[\s:|-]+\|?\s*$', lines[i + 1]) and '-' in lines[i + 1]:
            _flush_table()
            # 解析表头
            header = [c.strip() for c in stripped.strip('|').split('|')]
            # 跳过分隔行
            i += 2
            rows = [header]
            while i < n and '|' in lines[i].strip() and lines[i].strip():
                row = [c.strip() for c in lines[i].strip('|').split('|')]
                rows.append(row)
                i += 1
            table_buffer = rows
            _flush_table()
            continue

        # 标题
        hm = re.match(r'^(#{1,6})\s+(.*)$', stripped)
        if hm:
            _flush_table()
            level = min(len(hm.group(1)), 4)
            h = doc.add_heading(hm.group(2).strip(), level=level)
            for r in h.runs:
                _set_eastasia(r)
            i += 1
            continue

        # 引用块
        if stripped.startswith('>'):
            _flush_table()
            p = doc.add_paragraph()
            run = p.add_run(stripped.lstrip('>').strip())
            run.italic = True
            _set_eastasia(run)
            i += 1
            continue

        # 无序列表
        if re.match(r'^[-*]\s+', stripped):
            _flush_table()
            indent = len(raw) - len(raw.lstrip(' '))
            content = re.sub(r'^[-*]\s+', '', stripped)
            p = doc.add_paragraph(style='List Bullet' if indent < 4 else 'List Bullet 2')
            _add_inline(p, content)
            i += 1
            continue

        # 有序列表
        if re.match(r'^\d+[.)]\s+', stripped):
            _flush_table()
            content = re.sub(r'^\d+[.)]\s+', '', stripped)
            p = doc.add_paragraph(style='List Number')
            _add_inline(p, content)
            i += 1
            continue

        # 普通段落
        _flush_table()
        p = doc.add_paragraph()
        _add_inline(p, stripped)
        i += 1

    _flush_table()


def build_ir_step_docx(step_md_path: Path | str, output_path: Path | str,
                       title: str = "") -> str | None:
    """将一个 step 的 markdown 渲染为独立 DOCX。

    返回输出路径；输入不存在或渲染失败时返回 None。
    """
    step_md_path = Path(step_md_path)
    if not step_md_path.exists() or step_md_path.stat().st_size < 50:
        return None

    raw = step_md_path.read_text(encoding="utf-8", errors="ignore")
    text = sanitize_text(raw) if sanitize_text else raw

    doc = Document()
    _apply_base_style(doc)
    _render_markdown(doc, text, title=title or step_md_path.stem)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return str(output_path)


if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        _title = sys.argv[3] if len(sys.argv) > 3 else ""
        _out = sys.argv[2] if len(sys.argv) > 2 else (Path(sys.argv[1]).stem + ".docx")
        _built = build_ir_step_docx(sys.argv[1], _out, title=_title)
        print(_built or "FAILED")
